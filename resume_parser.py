"""
resume_parser.py
-----------------
Handles extracting raw text from uploaded resume files (PDF or DOCX).
"""

import io
import pdfplumber
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract all text from a PDF file given as bytes."""
    text_chunks = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_chunks.append(page_text)
    return "\n".join(text_chunks)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract all text from a DOCX file given as bytes."""
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also grab text inside tables (some resumes use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs)


def extract_resume_text(uploaded_file) -> str:
    """
    Given a Streamlit UploadedFile object, detect its type and extract text.
    Raises ValueError for unsupported file types.
    """
    filename = uploaded_file.name.lower()
    file_bytes = uploaded_file.getvalue()

    if filename.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
    elif filename.endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
    elif filename.endswith(".txt"):
        text = file_bytes.decode("utf-8", errors="ignore")
    else:
        raise ValueError(
            "Unsupported file type. Please upload a PDF, DOCX, or TXT file."
        )

    text = text.strip()
    if not text:
        raise ValueError(
            "Couldn't extract any text from this file. "
            "It might be a scanned/image-based PDF."
        )

    return text
